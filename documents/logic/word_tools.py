"""
Word Tools Module – Final Version

This module provides utilities for working with Microsoft Word (DOCX) files. It allows
creating documents from templates, setting core properties (e.g. author, title,
revision), and extracting core properties and comments.

It is designed to fail gracefully if the `python-docx` package is not installed.

The module also includes a self-test that can be enabled by setting the
module-level constant `SELFTEST` to 1. When the module is executed as
`__main__` with `SELFTEST` enabled, it runs a simple end-to-end test and
prints results to the console.
"""

from __future__ import annotations

# Selftest control: set to 1 to enable simple self test when run directly.
SELFTEST: int = 0  # 0 = disabled, 1 = enabled

import os
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Tuple
from xml.etree import ElementTree as ET

try:
    # python-docx is optional; import errors are deferred until use.
    from docx import Document  # type: ignore
    _HAVE_DOCX = True
except Exception:
    Document = None  # type: ignore
    _HAVE_DOCX = False


def _ensure_docx_available() -> None:
    """Raise an informative error if python-docx is not installed."""
    if not _HAVE_DOCX:
        raise RuntimeError(
            "python-docx is not installed. Please install it with:\n  pip install python-docx"
        )


def create_from_template(template_path: str, out_path: str, *, props: Dict[str, Any]) -> str:
    """
    Create a new DOCX file based on the given template and set core properties.

    Parameters
    ----------
    template_path: str
        Path to the DOCX template file.
    out_path: str
        Path where the new DOCX should be written.
    props: dict
        Dictionary of core properties to set, such as 'author', 'title', 'revision'.

    Returns
    -------
    str
        The absolute path to the written DOCX.
    """
    _ensure_docx_available()
    doc = Document(template_path)  # type: ignore[misc]
    _apply_core_properties(doc, props)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def set_core_properties(docx_path: str, *, props: Dict[str, Any]) -> None:
    """
    Update the core properties of an existing DOCX file.

    Parameters
    ----------
    docx_path: str
        Path to the DOCX file whose properties should be updated.
    props: dict
        Properties to set.
    """
    _ensure_docx_available()
    doc = Document(docx_path)  # type: ignore[misc]
    _apply_core_properties(doc, props)
    doc.save(docx_path)


def extract_core_and_comments(docx_path: str) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """
    Extract core properties and comments from a DOCX file.

    Parameters
    ----------
    docx_path: str
        Path to the DOCX file.

    Returns
    -------
    (dict, list)
        A tuple containing a dictionary of core properties and a list of comment dicts.
        Each comment dict has keys: 'author', 'date', and 'text'.
    """
    core: Dict[str, Any] = {}
    comments: List[Dict[str, Any]] = []
    if _HAVE_DOCX:
        try:
            d = Document(docx_path).core_properties  # type: ignore[misc]
            core = {
                "title": d.title or "",
                "subject": d.subject or "",
                "category": d.category or "",
                "keywords": d.keywords or "",
                "author": d.author or "",
                "last_modified_by": d.last_modified_by or "",
                "comments": d.comments or "",
                "created": d.created if d.created else None,
                "modified": d.modified if d.modified else None,
                "revision": int(d.revision) if d.revision else None,
            }
        except Exception:
            core = {}
    # Comments are stored in word/comments.xml
    try:
        with zipfile.ZipFile(docx_path) as zf:
            if "word/comments.xml" in zf.namelist():
                xml_bytes = zf.read("word/comments.xml")
                root = ET.fromstring(xml_bytes)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for c in root.findall("w:comment", ns):
                    author = c.attrib.get(f"{{{ns['w']}}}author", "") or c.attrib.get("author", "")
                    date_str = c.attrib.get(f"{{{ns['w']}}}date", "") or c.attrib.get("date", "")
                    dt = None
                    if date_str:
                        try:
                            dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
                        except Exception:
                            dt = None
                    texts: List[str] = []
                    for t in c.findall(".//w:t", ns):
                        texts.append(t.text or "")
                    comments.append({"author": author, "date": dt, "text": "".join(texts).strip()})
    except Exception:
        pass
    return core, comments


def _apply_core_properties(doc: Any, props: Dict[str, Any]) -> None:
    """
    Apply properties to a python-docx Document object.

    Parameters
    ----------
    doc: any
        An instance of python-docx Document.
    props: dict
        Key-value pairs for properties.
    """
    cp = doc.core_properties
    if "title" in props:
        cp.title = str(props["title"])
    if "subject" in props:
        cp.subject = str(props["subject"])
    if "category" in props:
        cp.category = str(props["category"])
    if "keywords" in props:
        cp.keywords = str(props["keywords"])
    if "author" in props:
        cp.author = str(props["author"])
    if "last_modified_by" in props:
        cp.last_modified_by = str(props["last_modified_by"])
    if "comments" in props:
        cp.comments = str(props["comments"])
    if "revision" in props and props["revision"] is not None:
        try:
            cp.revision = int(props["revision"])
        except Exception:
            pass


# Self test ---------------------------------------------------------------
if __name__ == "__main__" and SELFTEST:
    import tempfile
    print("[word_tools] Running selftest...")
    tmp_dir = tempfile.mkdtemp()
    template = os.path.join(tmp_dir, "template.docx")
    output = os.path.join(tmp_dir, "output.docx")
    # Create a basic document if python-docx is available
    if _HAVE_DOCX:
        doc = Document()  # type: ignore[misc]
        doc.add_paragraph("Hello, world.")
        doc.save(template)
        # Generate from template
        create_from_template(template, output, props={"author": "tester", "title": "Demo", "revision": 2})
        # Inject a fake comment
        xml = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="Max" w:date="2023-05-01T12:34:56Z">
    <w:p><w:r><w:t>This is a comment.</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""
        with zipfile.ZipFile(output, "a") as zf:
            zf.writestr("word/comments.xml", xml)
        core, comments = extract_core_and_comments(output)
        print("Core:", core)
        print("Comments:", comments)
    else:
        print("python-docx not installed; skipping selftest")
    print("[word_tools] Selftest complete.")